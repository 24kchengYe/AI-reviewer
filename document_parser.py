"""
文档解析模块
支持PDF和Word文档的文本提取
"""

import os
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document


class DocumentParser:
    """文档解析器"""

    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc']

    @staticmethod
    def is_supported(file_path):
        """检查文件是否支持"""
        ext = Path(file_path).suffix.lower()
        return ext in DocumentParser.SUPPORTED_EXTENSIONS

    @staticmethod
    def parse_pdf(file_path):
        """
        解析PDF文件
        优先使用pdfplumber，失败则使用PyPDF2
        """
        text = ""

        try:
            # 尝试使用pdfplumber（更好的文本提取效果）
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

            if text.strip():
                return text
        except Exception as e:
            print(f"pdfplumber解析失败: {e}, 尝试使用PyPDF2")

        try:
            # 备用方案：使用PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
            return text
        except Exception as e:
            raise Exception(f"PDF解析失败: {e}")

    @staticmethod
    def parse_docx(file_path):
        """解析Word文档"""
        try:
            doc = Document(file_path)
            text = ""

            # 提取段落文本
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text
        except Exception as e:
            raise Exception(f"Word文档解析失败: {e}")

    @staticmethod
    def parse(file_path):
        """
        解析文档（自动识别类型）
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentParser.parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentParser.parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")


def test_parser():
    """测试文档解析器"""
    import sys

    if len(sys.argv) < 2:
        print("用法: python document_parser.py <文件路径>")
        return

    file_path = sys.argv[1]

    try:
        text = DocumentParser.parse(file_path)
        print(f"成功解析文档，提取文本长度: {len(text)} 字符")
        print("\n前500字符预览:")
        print(text[:500])
    except Exception as e:
        print(f"解析失败: {e}")


if __name__ == "__main__":
    test_parser()
